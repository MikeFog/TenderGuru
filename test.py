from flask import Flask, jsonify, request
import os
import requests
from requests.auth import HTTPBasicAuth
from bs4 import BeautifulSoup
from pathlib import Path
import fitz  # PyMuPDF для работы с PDF
from docx import Document
import tiktoken
from openai import OpenAI

# Функция для извлечения текста из PDF
def extract_text_from_pdf(file_path):
    text = ""
    with fitz.open(file_path) as pdf:
        for page in pdf:
            text += page.get_text()
    return text

# Функция для извлечения текста из DOCX
def extract_text_from_docx(file_path):
    doc = Document(file_path)
    text = "\n".join([p.text for p in doc.paragraphs])
    return text

# Функция для извлечения текста (пример из предыдущих шагов)
def extract_text_from_file(file_path):
    if file_path.endswith(".pdf"):
        return extract_text_from_pdf(file_path)  # Используй свою функцию
    elif file_path.endswith(".docx"):
        return extract_text_from_docx(file_path)  # Используй свою функцию
    elif file_path.endswith(".txt"):
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    else:
        return ""

def split_text_into_chunks(text, max_tokens=8000):
    # Выбираем токенайзер
    encoding = tiktoken.encoding_for_model("gpt-4o-mini")

    # Преобразуем текст в токены
    tokens = encoding.encode(text)
    
    # Разделяем текст на части
    chunks = []
    while len(tokens) > max_tokens:
        chunk = tokens[:max_tokens]
        chunks.append(encoding.decode(chunk))
        tokens = tokens[max_tokens:]
    chunks.append(encoding.decode(tokens))  # Последний блок
    return chunks


app = Flask(__name__)

@app.route('/processUrl')
def greet():
    url = request.args.get('url', 'мир')

    base_url = "https://www.bicotender.ru"
    action = "/login/"
    login_url = requests.compat.urljoin(base_url, action)
    login_data = {
        "login": "client1399145",
        "password": "K1208660325"
    }

    # Создаем сессию
    session = requests.Session()

    # Логинимся
    response = session.post(login_url, data=login_data)
    response.raise_for_status()

    # URL защищенной страницы
    tender_url = url

    # Запрашиваем страницу тендера
    response = session.get(tender_url)
    response.raise_for_status()

    # Парсинг HTML
    soup = BeautifulSoup(response.text, 'html.parser')

    # Текущая рабочая директория
    base_dir = os.getcwd()

    # Каталог для сохранения файлов
    output_dir = os.path.join(base_dir, "temp_files")
    os.makedirs(output_dir, exist_ok=True)

    # Ищем все ссылки на документы
    files_path = []
    for link in soup.find_all('a', href=True):
        if "file_id" in link['href'] and "load" in link['href'] and "/browser/" not in link['href'] and "/all/" not in link['href']:
            doc_url = link['href']
            file_name = link.text.strip()  # Извлекаем текст внутри тега <a>
            
            # Полный путь для сохранения
            file_path = os.path.join(output_dir, file_name)
            
            doc_response = session.get(doc_url)
            if doc_response.status_code == 200:
                with open(file_path, 'wb') as file:
                    file.write(doc_response.content)
                #print(f"Файл сохранён: {file_path}")
                files_path.append(file_path)
            else:
                print(f"Ошибка при скачивании: {doc_response.status_code}")

    # Считываем promt из файла
    response = requests.get('https://onedrive.live.com/download?cid=7a4dfce935f80db4&resid=7A4DFCE935F80DB4!34319&authkey=!AGtblFPF02S2fOk', stream=True)

    # Проверяем, успешно ли выполнен запрос
    if response.status_code == 200:
        file_path = os.path.join(output_dir, 'promt.docx')
        with open(file_path, 'wb') as file:
            file.write(response.content)
        print(f'Файл успешно скачан и сохранен как {file_path}')
    else:
        print(f'Ошибка при скачивании файла: {response.status_code}')

    doc = Document(file_path)
    questions = [p.text for p in doc.paragraphs if p.text != ""]

    # Чтение текста из всех файлов
    files_text = {}

    for path in files_path:
        text = extract_text_from_file(path)
        file_name = Path(path).name
        if text:  # Пропускаем пустые файлы
            files_text[file_name] = text

    # Объединение текста
    combined_text = ""
    for file_name, text in files_text.items():
        combined_text += f"\n=== Начало текста из файла: {file_name} ===\n"
        combined_text += text
        combined_text += f"\n=== Конец текста из файла: {file_name} ===\n"


    # Разбиваем текст на блоки
    chunks = split_text_into_chunks(combined_text, max_tokens=8000)

    # Укажи свой API-ключ
    openai_key = os.getenv('OpenAI')
    client = OpenAI(api_key = openai_key)
    relevant_chunks = []

    # Обрабатываем каждый блок
    for i, chunk in enumerate(chunks, 1):
        print(f"Обрабатываем блок {i}/{len(chunks)}...")

        # Формируем запрос чтобы выделить только блоки, содержащие нужную информацию
        prompt = (
            f"Вот фрагмент текста:\n{chunk}\n\nОтветь содержит ли этот текст ответы на следующие вопросы. Твой ответ должен содержать только Да или Нет. Больше ничего отвечать, детализировать не надо:\n"
            )

        for idx, question in enumerate(questions, 1):
            prompt += f"{idx}. {question}\n"

        # Отправляем запрос в OpenAI
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "Ты сотрудник отдела закупок, который проводит анализ конкурсной документации."},
                {"role": "user", "content": prompt}
            ]
        )
        
        answer = response.choices[0].message.content
        if "да" in answer.lower():
            relevant_chunks.append(chunk)       

    #Сформировать запрос из релевантных фрагментов:
    combined_relevant_text = "\n".join(relevant_chunks)
    prompt = (
        f"Вот фрагмент текста:\n{combined_relevant_text}\n\nОтветь на следующие вопросы:\n"
        )

    for idx, question in enumerate(questions, 1):
        prompt += f"{idx}. {question}\n"

    # Отправляем запрос в OpenAI
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "Ты сотрудник отдела закупок, который проводит анализ конкурсной документации. Твоя задача — искать ключевую информацию"},
            {"role": "user", "content": prompt}
        ]
    )

    # Формируем сложный объект для возврата
    result = {
        "contest_name": "Конкурс №31231",
        "description": response.choices[0].message.content
    }

    return jsonify(result)

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 10000))  # Default to Render's port 10000
    app.run(host='0.0.0.0', port=port)